"""
DOCX Extractor for HoloLearn
Extracts text content from Word documents.
"""

from pathlib import Path
from typing import Dict, Any, Optional, List
from datetime import datetime
import json
import time
import re
import tempfile
import shutil
import zipfile

import sys
sys.path.append(str(Path(__file__).parent.parent))
from utils.configs import OUTPUT_DIR, LOGS_DIR
from utils.error_handler import ErrorHandler
from utils.text_cleaner import TextCleaner

try:
    from docx import Document
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

try:
    import win32com.client
    _WIN32COM_AVAILABLE = True
except ImportError:
    _WIN32COM_AVAILABLE = False


class DOCXExtractor:
    """Extract text from Word documents"""

    def __init__(self):
        if not _DOCX_AVAILABLE:
            raise ImportError(
                "python-docx not installed. Install with: pip install python-docx"
            )

        self.text_cleaner = TextCleaner()
        self.base_output_dir = OUTPUT_DIR
        self.base_logs_dir = LOGS_DIR

        self.base_output_dir.mkdir(parents=True, exist_ok=True)
        self.base_logs_dir.mkdir(parents=True, exist_ok=True)

    # ------------------------------------------------------------------ #
    #  Legacy .doc conversion                                              #
    # ------------------------------------------------------------------ #

    def _convert_doc_to_docx(self, doc_path: Path) -> Optional[Path]:
        """Convert a legacy .doc file to .docx via Word COM or LibreOffice."""
        tmp_dir = Path(tempfile.mkdtemp())
        docx_path = tmp_dir / (doc_path.stem + ".docx")

        if _WIN32COM_AVAILABLE:
            try:
                import pythoncom
                pythoncom.CoInitialize()
                word_app = win32com.client.Dispatch("Word.Application")
                word_app.Visible = False
                document = word_app.Documents.Open(str(doc_path.resolve()), ReadOnly=True)
                document.SaveAs2(str(docx_path.resolve()), FileFormat=16)  # 16 = wdFormatDocumentDefault
                document.Close()
                word_app.Quit()
                if docx_path.exists():
                    return docx_path
            except Exception:
                pass
            finally:
                try:
                    pythoncom.CoUninitialize()
                except Exception:
                    pass

        import subprocess
        for cmd in ("libreoffice", "soffice"):
            try:
                result = subprocess.run(
                    [cmd, "--headless", "--convert-to", "docx",
                     "--outdir", str(tmp_dir), str(doc_path.resolve())],
                    capture_output=True, timeout=60
                )
                if result.returncode == 0 and docx_path.exists():
                    return docx_path
            except (FileNotFoundError, subprocess.TimeoutExpired):
                continue

        shutil.rmtree(tmp_dir, ignore_errors=True)
        return None

    # ------------------------------------------------------------------ #
    #  Resource helpers                                                    #
    # ------------------------------------------------------------------ #

    def _create_resource_name(self, filename: str) -> str:
        name = Path(filename).stem.lower()
        name = re.sub(r'[^\w\s-]', '', name)
        name = re.sub(r'[-\s]+', '_', name).strip('_')
        return (name[:50] if len(name) > 50 else name) or "unnamed_resource"

    def _setup_resource_directories(self, resource_name: str, output_dir_override: Optional[Path] = None) -> tuple:
        resource_output_dir = Path(output_dir_override) if output_dir_override else self.base_output_dir / resource_name
        resource_logs_dir = self.base_logs_dir / resource_name
        resource_output_dir.mkdir(parents=True, exist_ok=True)
        resource_logs_dir.mkdir(parents=True, exist_ok=True)
        return resource_output_dir, resource_logs_dir

    # ------------------------------------------------------------------ #
    #  Structured extraction helpers                                       #
    # ------------------------------------------------------------------ #

    def _build_outline(self, doc) -> List[dict]:
        """Return flat list of heading entries [{level, title}] in document order."""
        outline = []
        for para in doc.paragraphs:
            if para.style.name.startswith('Heading') and para.text.strip():
                level_str = para.style.name.replace('Heading ', '')
                try:
                    level = int(level_str)
                except ValueError:
                    level = 1
                outline.append({"level": level, "title": para.text.strip()})
        return outline

    def _classify_paragraph_type(self, para) -> str:
        """Classify a paragraph's role for lecture generation."""
        style_name = para.style.name.lower()
        text = para.text.strip().lower()

        if style_name.startswith('heading'):
            return 'heading'
        if 'quote' in style_name or 'block' in style_name:
            return 'quote'
        if 'list' in style_name:
            return 'list_item'
        if re.match(r'^(definition|def)[.:]', text):
            return 'definition'
        if re.match(r'^(example|ex|e\.g\.)[.:]', text):
            return 'example'
        if re.match(r'^(note|important|warning|caution|tip)[.:]', text):
            return 'callout'
        return 'body'

    def _build_sections(self, doc) -> List[dict]:
        """Group paragraphs into sections delimited by headings."""
        sections: List[dict] = []
        current_title = ""
        current_type = "body"
        current_body: List[str] = []

        def _flush():
            if current_body:
                sections.append({
                    "title": current_title,
                    "body": "\n".join(current_body),
                    "type": current_type,
                    "source_location": {},
                })

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            if para.style.name.startswith('Heading'):
                _flush()
                current_body = []
                current_title = text
                level_str = para.style.name.replace('Heading ', '')
                try:
                    current_type = f"h{int(level_str)}"
                except ValueError:
                    current_type = "h1"
            else:
                para_type = self._classify_paragraph_type(para)
                current_body.append(f"[{para_type}] {text}" if para_type != 'body' else text)

        _flush()
        return sections

    def _extract_cross_refs(self, text: str) -> List[str]:
        """Detect in-document cross-references (see Figure X, refer to Chapter Y, etc.)."""
        patterns = [
            r'(?:see|refer to|as shown in|as described in|as mentioned in)\s+'
            r'(?:Figure|Fig\.|Table|Chapter|Section|Appendix)\s+[\w\d.]+',
            r'(?:Figure|Fig\.|Table|Chapter|Section|Appendix)\s+[\w\d.]+\s+'
            r'(?:shows|illustrates|demonstrates|describes)',
        ]
        refs = []
        for pattern in patterns:
            refs.extend(re.findall(pattern, text, re.IGNORECASE))
        return list(dict.fromkeys(r.strip() for r in refs))  # deduplicate, preserve order

    def _extract_comments(self, docx_path: Path) -> List[str]:
        """Extract reviewer comments from word/comments.xml inside the DOCX zip."""
        comments = []
        try:
            with zipfile.ZipFile(str(docx_path), 'r') as z:
                if 'word/comments.xml' not in z.namelist():
                    return comments
                xml_content = z.read('word/comments.xml').decode('utf-8', errors='ignore')
                comment_blocks = re.findall(
                    r'<w:comment\b[^>]*>(.*?)</w:comment>', xml_content, re.DOTALL
                )
                for block in comment_blocks:
                    texts = re.findall(r'<w:t(?:\s[^>]*)?>([^<]*)</w:t>', block)
                    comment_text = ' '.join(t.strip() for t in texts if t.strip())
                    if comment_text:
                        comments.append(comment_text)
        except Exception:
            pass
        return comments

    def _compute_quality_score(self, text: str, sections: List[dict]) -> float:
        word_score = min(len(text.split()) / 5000, 1.0)
        structure_score = min(len(sections) / 10, 1.0)
        return round(word_score * 0.6 + structure_score * 0.4, 2)

    # ------------------------------------------------------------------ #
    #  Main extraction                                                     #
    # ------------------------------------------------------------------ #

    def extract(self,
                docx_path: str,
                resource_id: Optional[str] = None,
                clean_text: bool = True,
                include_tables: bool = True,
                preserve_headings: bool = True,
                output_dir: Optional[str] = None) -> Dict[str, Any]:
        """
        Extract text from a Word document (.docx or .doc).

        Returns dict with: success, resource_name, resource_id, text_file,
        metadata_file, structured_file, output_dir, logs_dir, extracted_text,
        sections, outline, cross_refs, comments, content_quality_score, metadata.
        """
        start_time = time.time()
        docx_path = Path(docx_path)
        _tmp_dir_to_cleanup = None

        # Handle legacy .doc files
        if docx_path.suffix.lower() == ".doc":
            converted = self._convert_doc_to_docx(docx_path)
            if converted is None:
                resource_name = self._create_resource_name(docx_path.name)
                override = Path(output_dir) if output_dir else None
                out_dir, _ = self._setup_resource_directories(resource_name, output_dir_override=override)
                return self._create_error_result(
                    resource_name,
                    "Cannot convert .doc file: install Microsoft Office (win32com) or LibreOffice.",
                    out_dir, docx_path.name
                )
            _tmp_dir_to_cleanup = converted.parent
            docx_path = converted

        resource_name = self._create_resource_name(docx_path.name)
        override = Path(output_dir) if output_dir else None
        output_dir, logs_dir = self._setup_resource_directories(resource_name, output_dir_override=override)

        error_handler = ErrorHandler(f"docx_{resource_name}")
        error_handler.log_file = logs_dir / "extraction.log"
        error_handler.logger = error_handler._setup_logger()

        if not docx_path.exists():
            error_msg = f"DOCX file not found: {docx_path}"
            error_handler.log_error(FileNotFoundError(error_msg), context="Validating DOCX file",
                                    metadata={"path": str(docx_path)})
            return self._create_error_result(resource_name, error_msg, output_dir, docx_path.name)

        file_size = docx_path.stat().st_size
        file_size_mb = file_size / (1024 * 1024)

        error_handler.log_info(f"Starting DOCX extraction: {docx_path.name}",
                               metadata={"size_mb": f"{file_size_mb:.2f}",
                                         "resource_name": resource_name,
                                         "output_dir": str(output_dir)})

        try:
            doc = Document(docx_path)

            # --- flat text (backwards compatible) ---
            extracted_text = ""
            paragraph_count = 0
            heading_count = 0
            table_count = 0

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraph_count += 1
                    if preserve_headings and paragraph.style.name.startswith('Heading'):
                        heading_level = paragraph.style.name.replace('Heading ', '')
                        extracted_text += f"\n{'='*60}\n"
                        extracted_text += f"[HEADING {heading_level}] {paragraph.text}\n"
                        extracted_text += f"{'='*60}\n\n"
                        heading_count += 1
                    else:
                        extracted_text += paragraph.text + "\n"

            if include_tables and doc.tables:
                for table_num, table in enumerate(doc.tables, 1):
                    table_count += 1
                    extracted_text += f"\n--- TABLE {table_num} ---\n"
                    for row in table.rows:
                        row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_text:
                            extracted_text += " | ".join(row_text) + "\n"
                    extracted_text += "\n"

            if clean_text:
                extracted_text = self.text_cleaner.clean_text(
                    extracted_text, remove_urls=False, remove_emails=False, fix_spacing=True
                )
            extracted_text = self.text_cleaner.remove_duplicate_lines(extracted_text)

            # --- structured analysis ---
            outline = self._build_outline(doc)
            sections = self._build_sections(doc)
            cross_refs = self._extract_cross_refs(extracted_text)
            comments = self._extract_comments(docx_path)
            quality_score = self._compute_quality_score(extracted_text, sections)

            processing_time = time.time() - start_time
            core_props = doc.core_properties

            metadata = {
                "resource_name": resource_name,
                "resource_id": resource_id or resource_name,
                "filename": docx_path.name,
                "source_type": "docx",
                "upload_date": datetime.now().isoformat(),
                "extraction_timestamp": datetime.now().isoformat(),
                "file_size_bytes": file_size,
                "processing_time_seconds": round(processing_time, 2),
                "status": "success",
                "error_message": None,
                "paragraph_count": paragraph_count,
                "heading_count": heading_count,
                "table_count": table_count,
                "section_count": len(sections),
                "cross_ref_count": len(cross_refs),
                "comment_count": len(comments),
                "character_count": len(extracted_text),
                "included_tables": include_tables,
                "content_quality_score": quality_score,
                "document_properties": {
                    "title": core_props.title or "N/A",
                    "author": core_props.author or "N/A",
                    "subject": core_props.subject or "N/A",
                    "created": str(core_props.created) if core_props.created else "N/A",
                    "modified": str(core_props.modified) if core_props.modified else "N/A",
                },
            }

            text_file = output_dir / f"{resource_name}_text.txt"
            text_file.write_text(extracted_text, encoding='utf-8')
            metadata["extracted_text_path"] = str(text_file)

            metadata_file = output_dir / f"{resource_name}_metadata.json"
            metadata_file.write_text(json.dumps(metadata, indent=2), encoding='utf-8')

            structured = {
                "sections": sections,
                "outline": outline,
                "cross_refs": cross_refs,
                "comments": comments,
            }
            structured_file = output_dir / f"{resource_name}_structured.json"
            structured_file.write_text(json.dumps(structured, indent=2, ensure_ascii=False), encoding='utf-8')

            error_handler.log_success(f"DOCX extracted successfully: {docx_path.name}",
                                      metadata={"paragraphs": paragraph_count, "headings": heading_count,
                                                "sections": len(sections), "time": f"{processing_time:.2f}s"})

            return {
                "success": True,
                "resource_name": resource_name,
                "resource_id": resource_id or resource_name,
                "text_file": str(text_file),
                "metadata_file": str(metadata_file),
                "structured_file": str(structured_file),
                "output_dir": str(output_dir),
                "logs_dir": str(logs_dir),
                "extracted_text": extracted_text,
                "sections": sections,
                "outline": outline,
                "cross_refs": cross_refs,
                "comments": comments,
                "content_quality_score": quality_score,
                "metadata": metadata,
            }

        except Exception as e:
            processing_time = time.time() - start_time
            error_handler.log_error(e, context=f"Extracting DOCX: {docx_path.name}",
                                    metadata={"resource_name": resource_name})
            return self._create_error_result(resource_name, str(e), output_dir,
                                             docx_path.name, file_size, processing_time)
        finally:
            if _tmp_dir_to_cleanup:
                shutil.rmtree(_tmp_dir_to_cleanup, ignore_errors=True)

    def _create_error_result(self,
                             resource_name: str,
                             error_message: str,
                             output_dir: Path,
                             filename: str = "unknown",
                             file_size: int = 0,
                             processing_time: float = 0) -> Dict[str, Any]:
        metadata = {
            "resource_name": resource_name,
            "filename": filename,
            "source_type": "docx",
            "upload_date": datetime.now().isoformat(),
            "extraction_timestamp": datetime.now().isoformat(),
            "file_size_bytes": file_size,
            "processing_time_seconds": round(processing_time, 2),
            "status": "failed",
            "error_message": error_message,
        }
        metadata_file = output_dir / f"{resource_name}_metadata.json"
        metadata_file.write_text(json.dumps(metadata, indent=2), encoding='utf-8')
        return {
            "success": False,
            "resource_name": resource_name,
            "text_file": None,
            "metadata_file": str(metadata_file),
            "structured_file": None,
            "output_dir": str(output_dir),
            "extracted_text": "",
            "sections": [],
            "outline": [],
            "cross_refs": [],
            "comments": [],
            "content_quality_score": 0.0,
            "metadata": metadata,
            "error": error_message,
        }

    def extract_metadata_only(self, docx_path: str) -> Dict[str, Any]:
        """Extract only metadata without extracting text."""
        try:
            docx_path = Path(docx_path)
            doc = Document(docx_path)
            paragraph_count = len([p for p in doc.paragraphs if p.text.strip()])
            table_count = len(doc.tables)
            core_props = doc.core_properties
            return {
                "filename": docx_path.name,
                "resource_name": self._create_resource_name(docx_path.name),
                "paragraph_count": paragraph_count,
                "table_count": table_count,
                "file_size_bytes": docx_path.stat().st_size,
                "file_size_mb": round(docx_path.stat().st_size / (1024 * 1024), 2),
                "core_properties": {
                    "title": core_props.title or "N/A",
                    "author": core_props.author or "N/A",
                    "subject": core_props.subject or "N/A",
                    "created": str(core_props.created) if core_props.created else "N/A",
                    "modified": str(core_props.modified) if core_props.modified else "N/A",
                },
            }
        except Exception as e:
            ErrorHandler("docx_metadata").log_error(e, context=f"Extracting metadata from {docx_path}")
            return {}


# Example usage and testing
if __name__ == "__main__":
    from utils.file_picker import FilePicker

    print("=== Testing DOCX Extractor ===\n")

    extractor = DOCXExtractor()

    picker = FilePicker()
    print("Please select a Word document (.docx or .doc)...")
    test_docx = picker.pick_docx()
    picker.close()

    if test_docx:
        print(f"\n✓ Selected: {Path(test_docx).name}\n")

        print("1. Extracting metadata only...")
        metadata = extractor.extract_metadata_only(test_docx)
        print(f"   Resource name: {metadata.get('resource_name', 'N/A')}")
        print(f"   Paragraphs: {metadata.get('paragraph_count', 'N/A')}")
        print(f"   Tables: {metadata.get('table_count', 'N/A')}\n")

        print("2. Full extraction (with tables and headings)...")
        result = extractor.extract(
            docx_path=test_docx, clean_text=True,
            include_tables=True, preserve_headings=True
        )

        if result['success']:
            print(f"   ✓ Success!")
            print(f"   Sections detected: {len(result['sections'])}")
            print(f"   Outline entries: {len(result['outline'])}")
            print(f"   Cross-references: {len(result['cross_refs'])}")
            print(f"   Comments: {len(result['comments'])}")
            print(f"   Quality score: {result['content_quality_score']}")
            print(f"   Structured file: {result['structured_file']}")

            if result['outline']:
                print("\n   Outline (first 5 entries):")
                for entry in result['outline'][:5]:
                    indent = "  " * (entry['level'] - 1)
                    print(f"   {indent}H{entry['level']}: {entry['title']}")
        else:
            print(f"   ✗ Failed: {result['error']}")
    else:
        print("❌ No file selected")
        print("   The extractor is ready to use!")
